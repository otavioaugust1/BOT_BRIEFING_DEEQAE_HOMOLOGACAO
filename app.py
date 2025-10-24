# app.py
import importlib.util
import os
import sys

from flask import (Flask, jsonify, redirect, render_template, request,
                   send_file, send_from_directory, url_for)


# Configura√ß√µes para produ√ß√£o
app = Flask(__name__, static_folder='static', template_folder='templates')
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-key-123')

# Adiciona os diret√≥rios ao path para importar m√≥dulos locais
current_dir = os.path.dirname(os.path.abspath(__file__))
jobs_dir = os.path.join(current_dir, 'src', 'data_jobs', 'jobs')

# Adiciona ambos os diret√≥rios ao path
sys.path.append(current_dir)
sys.path.append(jobs_dir)

GERADOR_IMPORTADO = False


def carregar_gerador_relatorios():
    """Tenta carregar os m√≥dulos de gera√ß√£o de relat√≥rios"""
    global GERADOR_IMPORTADO, gerar_documento_briefing_completo, gerar_documento_briefing_simplificado

    try:
        # Tenta importar dos novos m√≥dulos
        from src.data_jobs.jobs.report_complete import \
            gerar_documento_briefing_completo
        from src.data_jobs.jobs.report_simplified import \
            gerar_documento_briefing_simplificado

        GERADOR_IMPORTADO = True
        print('‚úÖ M√≥dulos de relat√≥rio importados com sucesso')
        return True

    except ImportError as e:
        print(f'‚ùå Erro ao importar m√≥dulos: {e}')

        # Tenta m√©todo alternativo
        try:
            # Tenta importar report_complete
            spec_complete = importlib.util.spec_from_file_location(
                'report_complete', os.path.join(jobs_dir, 'report_complete.py')
            )
            report_complete = importlib.util.module_from_spec(spec_complete)
            spec_complete.loader.exec_module(report_complete)
            gerar_documento_briefing_completo = (
                report_complete.gerar_documento_briefing_completo
            )

            # Tenta importar report_simplified
            spec_simplified = importlib.util.spec_from_file_location(
                'report_simplified',
                os.path.join(jobs_dir, 'report_simplified.py'),
            )
            report_simplified = importlib.util.module_from_spec(
                spec_simplified
            )
            spec_simplified.loader.exec_module(report_simplified)
            gerar_documento_briefing_simplificado = (
                report_simplified.gerar_documento_briefing_simplificado
            )

            GERADOR_IMPORTADO = True
            print('‚úÖ M√≥dulos importados com sucesso (m√©todo alternativo)')
            return True

        except Exception as alt_e:
            print(f'‚ùå Falha na importa√ß√£o alternativa: {alt_e}')
            return False


# Tenta carregar os geradores
if not carregar_gerador_relatorios():
    print('üîÑ Usando fun√ß√µes de fallback...')
    GERADOR_IMPORTADO = False

    # Fun√ß√µes de fallback caso os m√≥dulos n√£o estejam dispon√≠veis
    def gerar_documento_briefing_simplificado(dados_selecao):
        from datetime import datetime

        from docx import Document

        # Cria um arquivo tempor√°rio de exemplo
        data_atual = datetime.now().strftime('%Y%m%d')
        nome_arquivo = f'{data_atual}_Briefing_Simplificado.docx'

        # Caminho para a pasta de downloads
        downloads_dir = os.path.join('static', 'downloads')
        os.makedirs(downloads_dir, exist_ok=True)
        caminho_arquivo = os.path.join(downloads_dir, nome_arquivo)

        # Cria um documento Word b√°sico
        doc = Document()
        doc.add_heading('BRIEFING SIMPLIFICADO - SISTEMA COQAE/DEEQAE', 0)
        doc.add_paragraph(f"Regi√£o: {dados_selecao.get('regiao', 'TODOS')}")
        doc.add_paragraph(f"UF: {dados_selecao.get('uf', 'TODOS')}")
        doc.add_paragraph(
            f"Macrorregi√£o: {dados_selecao.get('macro', 'TODOS')}"
        )
        doc.add_paragraph(
            f"Regi√£o de Sa√∫de: {dados_selecao.get('regiaoSaude', 'TODOS')}"
        )
        doc.add_paragraph(
            f"Munic√≠pio: {dados_selecao.get('municipio', 'TODOS')}"
        )
        doc.add_paragraph(f"Unidade: {dados_selecao.get('unidade', 'TODOS')}")
        doc.add_paragraph(
            f"Data de gera√ß√£o: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        )

        doc.save(caminho_arquivo)
        print(f'üìÑ Fallback: Documento simplificado salvo em {caminho_arquivo}')
        return f'static/downloads/{nome_arquivo}', nome_arquivo

    def gerar_documento_briefing_completo(dados_selecao):
        from datetime import datetime

        from docx import Document

        # Cria um arquivo tempor√°rio de exemplo
        data_atual = datetime.now().strftime('%Y%m%d')
        nome_arquivo = f'{data_atual}_Briefing_Completo.docx'

        # Caminho para a pasta de downloads
        downloads_dir = os.path.join('static', 'downloads')
        os.makedirs(downloads_dir, exist_ok=True)
        caminho_arquivo = os.path.join(downloads_dir, nome_arquivo)

        # Cria um documento Word b√°sico
        doc = Document()
        doc.add_heading('BRIEFING COMPLETO - SISTEMA COQAE/DEEQAE', 0)
        doc.add_paragraph(f"Regi√£o: {dados_selecao.get('regiao', 'TODOS')}")
        doc.add_paragraph(f"UF: {dados_selecao.get('uf', 'TODOS')}")
        doc.add_paragraph(
            f"Macrorregi√£o: {dados_selecao.get('macro', 'TODOS')}"
        )
        doc.add_paragraph(
            f"Regi√£o de Sa√∫de: {dados_selecao.get('regiaoSaude', 'TODOS')}"
        )
        doc.add_paragraph(
            f"Munic√≠pio: {dados_selecao.get('municipio', 'TODOS')}"
        )
        doc.add_paragraph(f"Unidade: {dados_selecao.get('unidade', 'TODOS')}")
        doc.add_paragraph(f'\nDETALHES ADICIONAIS:')
        doc.add_paragraph(f'- An√°lise aprofundada da regi√£o')
        doc.add_paragraph(f'- Indicadores detalhados')
        doc.add_paragraph(f'- Metadados completos')
        doc.add_paragraph(f'- Recomenda√ß√µes estrat√©gicas')
        doc.add_paragraph(
            f"Data de gera√ß√£o: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        )

        doc.save(caminho_arquivo)
        print(f'üìÑ Fallback: Documento completo salvo em {caminho_arquivo}')
        return f'static/downloads/{nome_arquivo}', nome_arquivo


# Garante que as pastas necess√°rias existam
def criar_pastas_necessarias():
    pastas_necessarias = [
        'static/downloads',
        'db/cnes',
        'db/ibge',
        'db/prefeitos',
        'db/secretario',
        'db/logs',  # Adicionando pasta de logs
        'templates',
    ]

    for pasta in pastas_necessarias:
        os.makedirs(pasta, exist_ok=True)


# Rota principal - mostra a p√°gina de login
@app.route('/')
def index():
    return render_template('index.html')


# Rota para servir arquivos do banco de dados
@app.route('/db/<path:filename>')
def serve_db_file(filename):
    try:
        return send_from_directory('db', filename)
    except FileNotFoundError:
        return jsonify({'error': 'Arquivo n√£o encontrado'}), 404


# Rota de login - processa o formul√°rio
@app.route('/login', methods=['POST'])
def login():
    username = request.form.get('username')
    password = request.form.get('password')

    # Valida√ß√£o simples para demonstra√ß√£o
    if username == 'COQAE' and password == '123456':
        return redirect(url_for('briefing_app'))
    else:
        return redirect(url_for('index'))


# Rota do aplicativo de briefing (ap√≥s login)
@app.route('/briefing')
def briefing_app():
    return render_template('briefing_app.html')


# Rota da gera√ß√£o do briefing-SIMPLIFICADO
@app.route('/gerar-briefing-simplificado', methods=['POST'])
def gerar_briefing_simplificado():
    try:
        dados_selecao = request.get_json()

        if not dados_selecao:
            return (
                jsonify({'success': False, 'error': 'Nenhum dado recebido'}),
                400,
            )

        print('Dados recebidos para briefing SIMPLIFICADO:', dados_selecao)

        # Usa o gerador real ou fallback
        caminho_arquivo, nome_arquivo = gerar_documento_briefing_simplificado(
            dados_selecao
        )

        print(f'Arquivo gerado: {nome_arquivo}')
        print(f'Caminho: {caminho_arquivo}')

        return jsonify(
            {
                'success': True,
                'arquivo': caminho_arquivo,
                'nomeArquivo': nome_arquivo,
            }
        )

    except Exception as e:
        print('Erro ao gerar briefing simplificado:', str(e))
        return (
            jsonify({'success': False, 'error': f'Erro interno: {str(e)}'}),
            500,
        )


# Rota da gera√ß√£o do briefing-COMPLETO
@app.route('/gerar-briefing-completo', methods=['POST'])
def gerar_briefing_completo():
    try:
        dados_selecao = request.get_json()

        if not dados_selecao:
            return (
                jsonify({'success': False, 'error': 'Nenhum dado recebido'}),
                400,
            )

        print('Dados recebidos para briefing COMPLETO:', dados_selecao)

        # Usa o gerador real ou fallback
        caminho_arquivo, nome_arquivo = gerar_documento_briefing_completo(
            dados_selecao
        )

        print(f'Arquivo gerado: {nome_arquivo}')
        print(f'Caminho: {caminho_arquivo}')

        return jsonify(
            {
                'success': True,
                'arquivo': caminho_arquivo,
                'nomeArquivo': nome_arquivo,
            }
        )

    except Exception as e:
        print('Erro ao gerar briefing completo:', str(e))
        return (
            jsonify({'success': False, 'error': f'Erro interno: {str(e)}'}),
            500,
        )


@app.route('/download/<path:filename>')
def download_file(filename):
    try:
        # Remove caracteres perigosos do filename
        safe_filename = os.path.basename(filename)
        safe_path = os.path.join('static', 'downloads', safe_filename)

        if not os.path.exists(safe_path):
            return jsonify({'error': 'Arquivo n√£o encontrado'}), 404

        return send_file(safe_path, as_attachment=True)

    except Exception as e:
        print(f'Erro no download: {e}')
        return jsonify({'error': 'Erro ao baixar arquivo'}), 500


# Rota de health check
@app.route('/health')
def health_check():
    return jsonify(
        {
            'status': 'ok',
            'gerador_importado': GERADOR_IMPORTADO,
            'pastas_existem': all(
                os.path.exists(pasta)
                for pasta in ['static', 'templates', 'db']
            ),
        }
    )


# Manipulador de erro 404
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint n√£o encontrado'}), 404


# Manipulador de erro 500
@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Erro interno do servidor'}), 500


if __name__ == '__main__':
    # Cria pastas necess√°rias ao iniciar
    criar_pastas_necessarias()

    print('=' * 50)
    print('SISTEMA BOT BRIEFING DEEQAE')
    print('=' * 50)
    print(
        f"Gerador importado: {'SIM' if GERADOR_IMPORTADO else 'N√ÉO (usando fallback)'}"
    )
    print('Pastas verificadas e criadas se necess√°rio')
    print('Servidor iniciando em: http://127.0.0.1:5000')
    print('=' * 50)

    app.run(debug=True, host='127.0.0.1', port=5000)
