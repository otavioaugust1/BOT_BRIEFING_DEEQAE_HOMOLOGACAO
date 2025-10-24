# ==============================================================================
#             CONFIGURA√á√ÉO DO RELATORIO COMPLETO - REGRA (10 PAGINAS)
# ==============================================================================

# report_complete.py
import os
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
# Importa configura√ß√µes compartilhadas
from report_configuration import (BASE_DIR, adicionar_cabecalho_com_logo,
                                  encontrar_hierarquia_completa,
                                  gerar_nome_prefeito, gerar_nome_secretario,
                                  verificar_arquivo_existente)

from src.data_jobs.jobs.processing_geral import (  # ‚úÖ NOVA FUN√á√ÉO ADICIONADA
    gerar_descricao_demografica, gerar_tabela_cnes_hab)
from src.data_jobs.jobs.processing_geral_2 import \
    gerar_tabela_cnes_srv  # ‚úÖ NOVA FUN√á√ÉO ADICIONADA


# Informa√ß√£o inicial
def add_info_paragraph_formatado(doc, titulo, valor):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run_titulo = p.add_run(f'‚Ä¢ {titulo}: ')
    run_titulo.bold = True
    run_valor = p.add_run(valor)


# Espa√ßo antes do rodap√©
def add_rodape_paragraph(doc, texto):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


# fonte_configura√ß√£o
def add_fonte_paragraph(doc, partes):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for texto, estilo in partes:
        run = p.add_run(texto)
        run.bold = estilo.get('bold', False)
        run.italic = estilo.get('italic', False)
        run.font.size = Pt(8)
        run.font.name = 'Arial'


# ==============================================================================
#                           RELATORIO COMPLETO
# ==============================================================================


def gerar_documento_briefing_completo(dados_selecao):
    """Gera documento completo com base nos dados de sele√ß√£o"""

    # Primeiro verifica se j√° existe um arquivo gerado hoje
    caminho_existente, nome_arquivo = verificar_arquivo_existente(
        dados_selecao, 'COMPLETO'
    )
    if caminho_existente:
        return caminho_existente, nome_arquivo

    # ENCONTRA a hierarquia completa baseada na sele√ß√£o
    dados_completos = encontrar_hierarquia_completa(dados_selecao)

    # Agora usa os dados completos para gerar o documento
    regiao = dados_completos.get('regiao', 'TODOS')
    uf = dados_completos.get('uf', 'TODOS')
    macro = dados_completos.get('macro', 'TODOS')
    regiao_saude = dados_completos.get('regiaoSaude', 'TODOS')
    municipio = dados_completos.get('municipio', 'TODOS')
    unidade = dados_completos.get('unidade', 'TODOS')

    print(f'üîÑ Gerando NOVO briefing COMPLETO com hierarquia completa:')
    print(f'   Regi√£o: {regiao}, UF: {uf}, Macro: {macro}')
    print(
        f'   Regi√£o Sa√∫de: {regiao_saude}, Munic√≠pio: {municipio}, Unidade: {unidade}'
    )

    # Caminho correto para a pasta static
    caminho_saida = os.path.join(
        BASE_DIR, '..', '..', '..', 'static', 'downloads', nome_arquivo
    )

    # Garante que o diret√≥rio existe
    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)

    # Cria o documento
    doc = Document()

    # Adiciona cabe√ßalho com logo
    adicionar_cabecalho_com_logo(doc)

    # T√≠tulo principal
    titulo = doc.add_heading('BRIEFING COMPLETO - SISTEMA COQAE/DEEQAE', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Se√ß√£o de dados da sele√ß√£o
    doc.add_heading('DADOS DA SELE√á√ÉO E CONTEXTUALIZA√á√ÉO', level=1)

    add_info_paragraph_formatado(doc, 'Regi√£o', regiao)
    add_info_paragraph_formatado(doc, 'UF', uf)
    add_info_paragraph_formatado(doc, 'Macrorregi√£o de Sa√∫de', macro)
    add_info_paragraph_formatado(doc, 'Regi√£o de Sa√∫de', regiao_saude)
    add_info_paragraph_formatado(doc, 'Munic√≠pio', municipio)
    add_info_paragraph_formatado(doc, 'Unidade', unidade)

    # Informa√ß√µes de gest√£o (condicional)
    if municipio != 'TODOS' and uf != 'TODOS':
        add_info_paragraph_formatado(
            doc, 'Prefeito(a)', gerar_nome_prefeito(municipio, uf)
        )
        add_info_paragraph_formatado(
            doc, 'Secret√°rio(a) de Sa√∫de', gerar_nome_secretario(municipio, uf)
        )

    # Linha divis√≥ria com espa√ßamento ajustado
    linha = doc.add_paragraph('_' * 50)
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linha.paragraph_format.space_before = Pt(0)
    linha.paragraph_format.space_after = Pt(0)
    linha.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Resumo Executivo Expandido
    doc.add_heading('RESUMO EXECUTIVO EXPANDIDO', level=1)
    doc.add_paragraph(
        'Este briefing, elaborado no contexto do programa "Agora Tem Especialistas", oferece uma vis√£o geral da situa√ß√£o de sa√∫de na regi√£o selecionada. A partir da consolida√ß√£o de dados p√∫blicos e oficiais, o documento re√∫ne informa√ß√µes relevantes sobre a estrutura, cobertura e desempenho dos servi√ßos de sa√∫de, contribuindo para o entendimento do cen√°rio local e subsidiando a atua√ß√£o dos profissionais especializados.'
    )

    # Dados Demogr√°ficos
    doc.add_heading('DADOS DEMOGR√ÅFICOS - IBGE', level=1)
    texto_demografico = gerar_descricao_demografica(dados_selecao)
    doc.add_paragraph(texto_demografico)

    # ‚úÖ NOVA SE√á√ÉO: HABILITA√á√ïES CNES
    doc.add_heading('HABILITA√á√ïES CNES - DISTRIBUI√á√ÉO HIER√ÅRQUICA', level=1)

    try:
        tabelas_cnes_hab = gerar_tabela_cnes_hab(dados_selecao)

        if tabelas_cnes_hab:
            for tabela_info in tabelas_cnes_hab:
                # T√≠tulo do tipo de habilita√ß√£o
                tipo_hab = tabela_info['tipo_habilitacao']
                doc.add_heading(tipo_hab.title(), level=2)

                # Cria a tabela no Word
                dados_tabela = tabela_info['dados']
                tabela = doc.add_table(rows=len(dados_tabela), cols=3)
                tabela.style = 'Table Grid'

                # Preenche a tabela
                for i, linha in enumerate(dados_tabela):
                    for j, valor in enumerate(linha):
                        tabela.cell(i, j).text = str(valor)

                # Formata√ß√£o da tabela
                for row in tabela.rows:
                    for cell in row.cells:
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Alinha a coluna QUANT √† direita
                for i in range(len(dados_tabela)):
                    cell = tabela.cell(i, 2)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                doc.add_paragraph()  # Espa√ßo entre tabelas
        else:
            doc.add_paragraph(
                'N√£o foram encontradas habilita√ß√µes CNES para os crit√©rios selecionados.'
            )

    except Exception as e:
        print(f'‚ö†Ô∏è Aviso: Erro ao gerar tabela CNES: {e}')
        doc.add_paragraph(
            'Dados de habilita√ß√µes CNES temporariamente indispon√≠veis.'
        )

    # ‚úÖ NOVA SE√á√ÉO: N√∫cleos de Gest√£o do Cuidado - NGC - CNES
    doc.add_heading(
        'N√öCLEOS DE GEST√ÉO DO CUIDADO - DISTRIBUI√á√ÉO HIER√ÅRQUICA', level=1
    )

    try:
        tabelas_cnes_hab = gerar_tabela_cnes_srv(dados_selecao)

        if tabelas_cnes_hab:
            for tabela_info in tabelas_cnes_hab:
                # T√≠tulo do tipo de habilita√ß√£o
                tipo_hab = tabela_info['tipo_habilitacao']
                doc.add_heading(tipo_hab.title(), level=2)

                # Cria a tabela no Word
                dados_tabela = tabela_info['dados']
                tabela = doc.add_table(rows=len(dados_tabela), cols=3)
                tabela.style = 'Table Grid'

                # Preenche a tabela
                for i, linha in enumerate(dados_tabela):
                    for j, valor in enumerate(linha):
                        tabela.cell(i, j).text = str(valor)

                # Formata√ß√£o da tabela
                for row in tabela.rows:
                    for cell in row.cells:
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Alinha a coluna QUANT √† direita
                for i in range(len(dados_tabela)):
                    cell = tabela.cell(i, 2)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                doc.add_paragraph()  # Espa√ßo entre tabelas
        else:
            doc.add_paragraph(
                'N√£o foram encontrada servi√ßo NGC CNES para os crit√©rios selecionados.'
            )

    except Exception as e:
        print(f'‚ö†Ô∏è Aviso: Erro ao gerar tabela CNES: {e}')
        doc.add_paragraph(
            'Dados de servi√ßo CNES temporariamente indispon√≠veis.'
        )

    # Recomenda√ß√µes Estrat√©gicas
    doc.add_heading('RECOMENDA√á√ïES ESTRAT√âGICAS', level=1)

    recomendacoes = [
        'Ampliar a cobertura de aten√ß√£o prim√°ria em √°reas de maior vulnerabilidade',
        'Fortalecer a rede de aten√ß√£o psicossocial na regi√£o',
        'Implementar programa de telemedicina para especialidades de dif√≠cil acesso',
        'Capacitar profissionais para gest√£o de cr√¥nicos',
        'Otimizar a distribui√ß√£o de medicamentos essenciais',
        'Fortalecer vigil√¢ncia em sa√∫de com foco em doen√ßas emergentes',
    ]

    for i, recomendacao in enumerate(recomendacoes, 1):
        doc.add_paragraph(f'{i}. {recomendacao}')

    # Metadados Expandidos
    doc.add_heading('METADADOS EXPANDIDOS', level=1)
    doc.add_paragraph(
        f'‚Ä¢ Data de Gera√ß√£o: {datetime.now().strftime("%d/%m/%Y √†s %H:%M:%S")}'
    )
    doc.add_paragraph(f'‚Ä¢ Tipo: Briefing Completo')
    doc.add_paragraph(f'‚Ä¢ Arquivo: {nome_arquivo}')
    doc.add_paragraph(f'‚Ä¢ Sistema: COQAE/DEEQAE - Minist√©rio da Sa√∫de')
    doc.add_paragraph(f'‚Ä¢ Per√≠odo de An√°lise: 2023-2025')

    # Fontes:
    doc.add_heading('FONTE DOS DADOS', level=1)

    add_fonte_paragraph(
        doc,
        [
            ('--- CONASEMS. ', {}),
            ('Rede COSEMS ‚Äì Dados. ', {'bold': True}),
            (
                'Dispon√≠vel em: https://portal.conasems.org.br/rede-cosems/dados. Acesso em: ',
                {},
            ),
            (datetime.now().strftime('%d/%m/%Y √†s %H:%M:%S'), {}),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. Minist√©rio da Sa√∫de. ', {}),
            ('Programa Agora Tem Especialistas ‚Äì InvestSUS. ', {'bold': True}),
            (
                'Dispon√≠vel em: https://investsuspaineis.saude.gov.br/extensions/CGIN_PMAE/CGIN_PMAE.html#. Acesso em: ',
                {},
            ),
            (datetime.now().strftime('%d/%m/%Y √†s %H:%M:%S'), {}),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. Minist√©rio da Sa√∫de. ', {}),
            ('Painel PNRF ‚Äì DRAC. ', {'bold': True}),
            (
                'Dispon√≠vel em: https://controleavaliacao.saude.gov.br/painel/pnrf.',
                {},
            ),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. Minist√©rio da Sa√∫de. ', {}),
            (
                'Macrorregi√µes e Regi√µes de Sa√∫de ‚Äì SEIDIGI/DEMAS. ',
                {'bold': True},
            ),
            (
                'Dispon√≠vel em: https://infoms.saude.gov.br/extensions/SEIDIGI_DEMAS_MACRORREGIOES/SEIDIGI_DEMAS_MACRORREGIOES.html. Acesso em: ',
                {},
            ),
            (datetime.now().strftime('%d/%m/%Y √†s %H:%M:%S'), {}),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. Minist√©rio da Sa√∫de. ', {}),
            ('Rede Nacional de Dados em Sa√∫de ‚Äì RNDS. ', {'bold': True}),
            ('Dispon√≠vel em: https://rnds-guia.saude.gov.br/.', {}),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. ', {}),
            ('DATASUS. ', {'italic': True}),
            (
                'Cadastro Nacional de Estabelecimentos de Sa√∫de ‚Äì CNES. ',
                {'bold': True},
            ),
            ('Dispon√≠vel em: https://cnes.datasus.gov.br/. Acesso em: ', {}),
            (datetime.now().strftime('%d/%m/%Y √†s %H:%M:%S'), {}),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. ', {}),
            ('DATASUS. ', {'italic': True}),
            ('Sistema de Informa√ß√µes Ambulatoriais ‚Äì SIA. ', {'bold': True}),
            (
                'Dispon√≠vel em: http://sia.datasus.gov.br/principal/index.php.',
                {},
            ),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. ', {}),
            ('DATASUS. ', {'italic': True}),
            ('Sistema de Informa√ß√µes Hospitalares ‚Äì SIH. ', {'bold': True}),
            ('Dispon√≠vel em: http://sih.datasus.gov.br/.', {}),
        ],
    )

    add_fonte_paragraph(
        doc,
        [
            ('--- BRASIL. ', {}),
            (
                'Instituto Brasileiro de Geografia e Estat√≠stica ‚Äì IBGE. ',
                {'bold': True},
            ),
            (
                'Portal IBGE. Dispon√≠vel em: https://www.ibge.gov.br/. Acesso em: ',
                {},
            ),
            (datetime.now().strftime('%d/%m/%Y √†s %H:%M:%S'), {}),
        ],
    )

    # Espa√ßo antes do rodap√©
    doc.add_paragraph('\n' * 3)

    # Linha divis√≥ria
    linha = doc.add_paragraph('_' * 50)
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linha.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    linha.paragraph_format.space_before = Pt(0)
    linha.paragraph_format.space_after = Pt(0)

    # Rodap√© visual com espa√ßamento ajustado
    add_rodape_paragraph(
        doc, 'Sistema de Gera√ß√£o Autom√°tica de Briefing - COQAE/DEEQAE'
    )
    add_rodape_paragraph(doc, 'Minist√©rio da Sa√∫de - Brasil')
    add_rodape_paragraph(
        doc, 'Departamento de Economia da Sa√∫de e Gest√£o Estrat√©gica'
    )
    add_rodape_paragraph(
        doc,
        f'Documento gerado automaticamente em: {datetime.now().strftime("%d/%m/%Y √†s %H:%M:%S")}',
    )

    # Salva o documento
    doc.save(caminho_saida)
    print(f'‚úÖ NOVO documento completo salvo em: {caminho_saida}')

    # Retorna caminho relativo para o Flask
    caminho_relativo = os.path.join('static', 'downloads', nome_arquivo)
    return caminho_relativo, nome_arquivo
