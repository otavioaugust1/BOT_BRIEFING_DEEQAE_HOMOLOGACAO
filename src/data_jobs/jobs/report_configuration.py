# ==============================================================================
#                   CONFIGURAÇÃO INICIAIS DO RELATORIO
# ==============================================================================


# report_configuration.py
import csv
import os
from datetime import datetime

import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Define o diretório base
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------
#         --- INICIO FUNÇÕES AUXILIARES DE CÓDIGO ---
# ---------------------------------------------------------------

# --- NOVA FUNÇÃO AUXILIAR DE NORMALIZAÇÃO DE STRING ---
def normalizar_string(texto):
    """Remove acentos, caracteres especiais (mantendo espaços) e converte para maiúsculas."""
    if not texto:
        return ''

    texto = str(texto).upper().strip()

    # Mapeamento para remover acentos e caracteres especiais
    subs = {
        'ÁÀÂÃÄ': 'A',
        'ÉÈÊË': 'E',
        'ÍÌÎÏ': 'I',
        'ÓÒÔÕÖ': 'O',
        'ÚÙÛÜ': 'U',
        'Ç': 'C',
        'Ñ': 'N',
        '&': 'E',
        '-': ' ',
        '/': ' ',
    }

    for chars, rep in subs.items():
        for char in chars:
            texto = texto.replace(char, rep)

    # Remove múltiplos espaços e transforma em um único espaço
    import re

    texto = re.sub(r'\s+', ' ', texto).strip()

    return texto


# --- FUNÇÕES AUXILIARES DE CÓDIGO (AJUSTADAS) ---
def extrair_codigo_regiao(regiao):
    """Extrai código da região"""
    if regiao == 'TODOS' or not regiao:
        return 'TD'
    mapeamento = {
        'NORDESTE': 'NE',
        'NORTE': 'NO',
        'CENTRO-OESTE': 'CO',
        'SUDESTE': 'SE',
        'SUL': 'SU',
    }
    return mapeamento.get(normalizar_string(regiao), regiao[:2].upper())


# --- FUNÇÕES EXTRAIR UF (SIGLA) ----
def extrair_codigo_uf(uf):
    """Extrai código da UF (SIGLA de 2 letras)"""
    if uf == 'TODOS' or not uf:
        return 'TD'

    # Mapeamento do nome completo NORMALIZADO para a sigla (ajuda na normalização para Parquet)
    mapeamento_sigla = {
        'ACRE': 'AC',
        'ALAGOAS': 'AL',
        'AMAPA': 'AP',
        'AMAZONAS': 'AM',
        'BAHIA': 'BA',
        'CEARA': 'CE',
        'DISTRITO FEDERAL': 'DF',
        'ESPIRITO SANTO': 'ES',
        'GOIAS': 'GO',
        'MARANHAO': 'MA',
        'MATO GROSSO': 'MT',
        'MATO GROSSO DO SUL': 'MS',
        'MINAS GERAIS': 'MG',
        'PARA': 'PA',
        'PARAIBA': 'PB',
        'PARANA': 'PR',
        'PERNAMBUCO': 'PE',
        'PIAUÍ': 'PI',
        'RIO DE JANEIRO': 'RJ',
        'RIO GRANDE DO NORTE': 'RN',
        'RIO GRANDE DO SUL': 'RS',
        'RONDONIA': 'RO',
        'RORAIMA': 'RR',
        'SANTA CATARINA': 'SC',
        'SAO PAULO': 'SP',
        'SERGIPE': 'SE',
        'TOCANTINS': 'TO',
    }

    # Normaliza o nome da UF de entrada (Ex: SAO PAULO)
    uf_normalizado = normalizar_string(uf)

    # Tenta usar o mapeamento.
    uf_sigla = mapeamento_sigla.get(uf_normalizado, uf_normalizado)

    # Se ainda for o nome completo (mas não encontrou no mapeamento), pega as 2 primeiras
    if len(uf_sigla) > 2:
        return uf_sigla[:2]

    return uf_sigla

# --- FUNÇÕES EXTRAIR MUNICIPIO (IBGE - TEXTO) ---
def extrair_codigo_macro(macro):
    """Extrai código IBGE do município"""
    if macro == 'TODOS' or not macro:
        return 'TD'
    # Se contém código entre parênteses, extrai
    if '(' in macro and ')' in macro:
        start = macro.find('(') + 1
        end = macro.find(')')
        return macro[start:end]
    return normalizar_string(macro)[:4] if macro else 'TD'


# --- FUNÇÕES EXTRAIR MUNICIPIO (IBGE - TEXTO) ---
def extrair_codigo_regiao_saude(regiao_saude):
    """Extrai código IBGE do município"""
    if regiao_saude == 'TODOS' or not regiao_saude:
        return 'TD'
    # Se contém código entre parênteses, extrai
    if '(' in regiao_saude and ')' in regiao_saude:
        start = regiao_saude.find('(') + 1
        end = regiao_saude.find(')')
        return regiao_saude[start:end]
    return normalizar_string(regiao_saude)[:3] if regiao_saude else 'TD'


# --- FUNÇÕES EXTRAIR MUNICIPIO (IBGE - TEXTO) ---
def extrair_codigo_ibge(municipio):
    """Extrai código IBGE do município"""
    if municipio == 'TODOS' or not municipio:
        return 'TD'
    # Se contém código entre parênteses, extrai
    if '(' in municipio and ')' in municipio:
        start = municipio.find('(') + 1
        end = municipio.find(')')
        return municipio[start:end]
    return normalizar_string(municipio)[:7] if municipio else 'TD'


# --- FUNÇÕES EXTRAIR UNIDADE - CNES ---
def extrair_codigo_cnes(unidade):
    """Extrai código CNES da unidade"""
    if unidade == 'TODOS' or not unidade:
        return 'TD'
    # Se contém código entre parênteses, extrai
    if '(' in unidade and ')' in unidade:
        start = unidade.find('(') + 1
        end = unidade.find(')')
        return unidade[start:end]
    return str(unidade)[:6] if unidade else 'TD'


# ---------------------------------------------------------------
#         --- FIM FUNÇÕES AUXILIARES DE CÓDIGO ---
# ---------------------------------------------------------------

# --- FUÇÃO CARREGAMENTO DE CSV - CONSULTA HIERARQUIA ----
def carregar_dados_csv():
    """Carrega os dados do CSV para consulta de hierarquia"""
    try:
        csv_path = os.path.join(
            BASE_DIR, '..', '..', '..', 'db', 'cnes', 'unidade_cnes.csv'
        )

        dados = []
        with open(csv_path, 'r', encoding='utf-8') as file:
            reader = csv.DictReader(file, delimiter=';')
            for row in reader:
                dados.append(
                    {
                        'regiao': row.get('REGIAO', ''),
                        'uf': row.get('UF_DESC', ''),
                        'macro': row.get('CO_MACROREGIAO_SAUDE', ''),
                        'regiaoSaude': row.get('CO_REGIAO_SAUDE', ''),
                        'municipio': row.get('MUNICIPIO', ''),
                        'unidade': f"{row.get('NOME_FANTASIA', '')} - {row.get('CNES', '')}",
                    }
                )
        return dados
    except Exception as e:
        print(f'⚠️ Aviso: Não foi possível carregar dados do CSV: {e}')
        return []


# ---- ENCONTRAR HIERARQUIA COMPLETA ----
def encontrar_hierarquia_completa(dados_selecao):
    """Encontra a hierarquia completa baseada nos dados selecionados"""
    dados_csv = carregar_dados_csv()
    if not dados_csv:
        return dados_selecao

    uf = dados_selecao.get('uf', 'TODOS')
    municipio = dados_selecao.get('municipio', 'TODOS')

    # Se temos um município específico, busca seus dados completos
    if municipio != 'TODOS' and uf != 'TODOS':
        for dado in dados_csv:
            if dado['municipio'] == municipio and dado['uf'] == uf:
                return {
                    'regiao': dado['regiao'],
                    'uf': dado['uf'],
                    'macro': dado['macro'],
                    'regiaoSaude': dado['regiaoSaude'],
                    'municipio': dado['municipio'],
                    'unidade': dados_selecao.get('unidade', 'TODOS'),
                }

    # Se só temos UF, busca a região dessa UF
    elif uf != 'TODOS':
        for dado in dados_csv:
            if dado['uf'] == uf:
                return {
                    'regiao': dado['regiao'],
                    'uf': dado['uf'],
                    'macro': dados_selecao.get('macro', 'TODOS'),
                    'regiaoSaude': dados_selecao.get('regiaoSaude', 'TODOS'),
                    'municipio': dados_selecao.get('municipio', 'TODOS'),
                    'unidade': dados_selecao.get('unidade', 'TODOS'),
                }

    # Se não encontrou, retorna os dados originais
    return dados_selecao


# ---- VERIFICAR ARQUIVO EXISTENTE FEITO NO DIA ----
def verificar_arquivo_existente(dados_selecao, tipo):
    """Verifica se já existe um arquivo gerado hoje com os mesmos parâmetros"""
    # PRIMEIRO: Encontra a hierarquia completa para gerar nome correto
    dados_completos = encontrar_hierarquia_completa(dados_selecao)

    # Extrai dados da hierarquia completa
    regiao = dados_completos.get('regiao', 'TODOS')
    uf = dados_completos.get('uf', 'TODOS')
    macro = dados_completos.get('macro', 'TODOS')
    regiao_saude = dados_completos.get('regiaoSaude', 'TODOS')
    municipio = dados_completos.get('municipio', 'TODOS')
    unidade = dados_completos.get('unidade', 'TODOS')

    # Gera nome do arquivo conforme regra
    data_atual = datetime.now().strftime('%Y%m%d')

    # Extrai códigos
    cod_uf = extrair_codigo_uf(uf)
    cod_ibge = extrair_codigo_ibge(municipio)
    cod_cnes = extrair_codigo_cnes(unidade)
    cod_regiao = extrair_codigo_regiao(regiao)
    cod_macro = extrair_codigo_macro(macro)
    cod_regiao_saude = extrair_codigo_regiao_saude(regiao_saude)

    if tipo == 'SIMPLIFICADO':
        nome_arquivo = f'{data_atual}_{cod_regiao}_{cod_uf}_{cod_macro}_{cod_regiao_saude}_{cod_ibge}_{cod_cnes}-RELATORIO_SIMPLIFICADO.docx'
    else:
        nome_arquivo = f'{data_atual}_{cod_regiao}_{cod_uf}_{cod_macro}_{cod_regiao_saude}_{cod_ibge}_{cod_cnes}-RELATORIO_COMPLETO.docx'

    # Remove caracteres inválidos
    nome_arquivo = ''.join(
        c for c in nome_arquivo if c.isalnum() or c in ('-', '_', '.')
    ).replace(' ', '_')

    # Caminho completo do arquivo
    caminho_arquivo = os.path.join(
        BASE_DIR, '..', '..', '..', 'static', 'downloads', nome_arquivo
    )

    # Verifica se o arquivo existe e foi criado hoje
    if os.path.exists(caminho_arquivo):
        data_criacao = datetime.fromtimestamp(
            os.path.getctime(caminho_arquivo)
        )
        data_hoje = datetime.now()

        # Se o arquivo foi criado hoje, retorna o caminho existente
        if data_criacao.date() == data_hoje.date():
            print(
                f'✅ Arquivo {tipo} já existe e foi gerado hoje: {nome_arquivo}'
            )
            caminho_relativo = os.path.join(
                'static', 'downloads', nome_arquivo
            )
            return caminho_relativo, nome_arquivo

    # Se não existe ou não foi criado hoje, retorna None
    return None, nome_arquivo


# ---- VERIFICAR PREFEITOS NA BASE ----
def carregar_prefeitos():
    """Carrega dados dos prefeitos do arquivo Parquet"""
    try:
        prefeitos_path = os.path.join(
            BASE_DIR,
            '..',
            '..',
            '..',
            'db',
            'prefeitos',
            'prefeitos_2024_BRASIL.parquet',
        )
        df_prefeitos = pd.read_parquet(prefeitos_path)
        return df_prefeitos
    except Exception as e:
        print(f'⚠️ Aviso: Não foi possível carregar dados dos prefeitos: {e}')
        return pd.DataFrame()


# ---- VERIFICAR SECRETARIO MUNICIPAL NA BASE ---
def carregar_secretarios():
    """Carrega dados dos secretários de saúde do arquivo Parquet"""
    try:
        secretarios_path = os.path.join(
            BASE_DIR,
            '..',
            '..',
            '..',
            'db',
            'secretario',
            'secretarios_municipais.parquet',
        )
        df_secretarios = pd.read_parquet(secretarios_path)
        return df_secretarios
    except Exception as e:
        print(
            f'⚠️ Aviso: Não foi possível carregar dados dos secretários: {e}'
        )
        return pd.DataFrame()


# --- CARREGAMENTO DO NOME DO PREFEITO NO RELATORIO ----
def gerar_nome_prefeito(municipio, uf):
    """Gera informações do prefeito baseado no município e UF"""
    if municipio == 'TODOS' or uf == 'TODOS':
        return 'Informação não disponível para seleção ampla'

    df_prefeitos = carregar_prefeitos()
    if df_prefeitos.empty:
        return 'Dados de prefeitos não disponíveis'

    # 1. NORMALIZAÇÃO DE DADOS DE ENTRADA
    # Normaliza o nome do município (TAUBATÉ -> TAUBATE)
    municipio_normalizado = normalizar_string(municipio)

    # Normaliza a UF (SAO PAULO -> SP)
    uf_sigla = extrair_codigo_uf(uf)

    # 2. NORMALIZAÇÃO DO DATAFRAME
    # Aplica a mesma normalização (sem acentos, maiúsculas) na coluna de município do DF
    df_prefeitos['NM_UE_NORMALIZED'] = df_prefeitos['NM_UE'].apply(
        normalizar_string
    )
    # A UF no Parquet (SG_UF) geralmente já é a sigla, mas garantimos que está em MAIÚSCULAS
    df_prefeitos['SG_UF_UPPER'] = df_prefeitos['SG_UF'].astype(str).str.upper()

    # 3. FILTRAGEM
    resultado = df_prefeitos[
        (df_prefeitos['NM_UE_NORMALIZED'] == municipio_normalizado)
        & (df_prefeitos['SG_UF_UPPER'] == uf_sigla)
    ]

    if not resultado.empty:
        prefeito = resultado.iloc[0]
        nome = prefeito.get('NM_URNA_CANDIDATO', 'N/A')
        partido = prefeito.get('SG_PARTIDO', 'N/A')
        coligacao = prefeito.get('NM_COLIGACAO', 'N/A')
        composicao = prefeito.get('DS_COMPOSICAO_COLIGACAO', 'N/A')

        info_prefeito = f'{nome} - ({partido})'
        if coligacao and coligacao != 'N/A':
            info_prefeito += f'\nColigação: {coligacao}'
        if composicao and composicao != 'N/A':
            info_prefeito += f'\nComposição: {composicao}'

        return info_prefeito
    else:
        print(
            f"DEBUG: Falha na busca por: MUNICÍPIO='{municipio_normalizado}', UF='{uf_sigla}'"
        )
        return f'Prefeito não encontrado para {municipio}/{uf}'


# --- CARREGAMENTO DO NOME DO SECRETARIO MUNICIPAL NO RELATORIO ---
def gerar_nome_secretario(municipio, uf):
    """Gera informações do secretário de saúde baseado no município e UF"""
    if municipio == 'TODOS' or uf == 'TODOS':
        return 'Informação não disponível para seleção ampla'

    df_secretarios = carregar_secretarios()
    if df_secretarios.empty:
        return 'Dados de secretários não disponíveis'

    # Normaliza os dados de entrada
    municipio_normalizado = normalizar_string(municipio)
    uf_sigla = extrair_codigo_uf(uf)

    # Normaliza as colunas de comparação no DataFrame
    df_secretarios['Municipio_NORMALIZED'] = df_secretarios['Municipio'].apply(
        normalizar_string
    )
    df_secretarios['UF_NORMALIZED'] = df_secretarios['UF'].apply(
        normalizar_string
    )

    # 1. Tenta buscar usando o município normalizado e a sigla (SP)
    resultado = df_secretarios[
        (df_secretarios['Municipio_NORMALIZED'] == municipio_normalizado)
        & (df_secretarios['UF_NORMALIZED'] == uf_sigla)
    ]

    # 2. Se não encontrou com a sigla, tenta usando o município normalizado e o nome do estado normalizado (SAO PAULO)
    if resultado.empty:
        uf_normalizado = normalizar_string(uf)
        resultado = df_secretarios[
            (df_secretarios['Municipio_NORMALIZED'] == municipio_normalizado)
            & (df_secretarios['UF_NORMALIZED'] == uf_normalizado)
        ]

    if not resultado.empty:
        secretario = resultado.iloc[0]
        nome = secretario.get('Nome', 'N/A')
        endereco = secretario.get('Endereço', 'N/A')
        cep = secretario.get('CEP', 'N/A')

        info_secretario = (
            f'{nome}' if nome != 'N/A' else 'Secretário não identificado'
        )

        if endereco and endereco != 'N/A':
            info_secretario += f'\nEndereço: {endereco}'
        if cep and cep != 'N/A':
            info_secretario += f'\nCEP: {cep}'

        return info_secretario
    else:
        return f'Secretário de saúde não encontrado para {municipio}/{uf}'


# --- LOGO - TIMBE - MARCA D'AGUA NO RELATORIO ---
def adicionar_cabecalho_com_logo(doc):
    """Adiciona cabeçalho com logo ao documento"""
    try:
        # Tenta adicionar o logo
        logo_path = os.path.join(
            BASE_DIR, '..', '..', '..', 'static', 'img', 'logo.png'
        )
        if os.path.exists(logo_path):
            # Adiciona uma seção com o logo
            section = doc.sections[0]
            header = section.header
            paragraph = (
                header.paragraphs[0]
                if header.paragraphs
                else header.add_paragraph()
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(
                'logo.png', width=Inches(6.5)
            )  # largura padrão da página A4
    except Exception as e:
        print(f'Aviso: Não foi possível adicionar o logo: {e}')


#
