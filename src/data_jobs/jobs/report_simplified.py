# ==============================================================================
#             CONFIGURAÇÃO DO RELATORIO SIMPLIFICADO - REGRA (1PAGINA)
# ==============================================================================


# report_simplified.py
import os
from datetime import datetime

from docx import Document
# Importa configurações compartilhadas
from report_configuration import (BASE_DIR, adicionar_cabecalho_com_logo,
                                  encontrar_hierarquia_completa,
                                  verificar_arquivo_existente)


def gerar_documento_briefing_simplificado(dados_selecao):
    """Gera documento simplificado com base nos dados de seleção"""

    # Primeiro verifica se já existe um arquivo gerado hoje
    caminho_existente, nome_arquivo = verificar_arquivo_existente(
        dados_selecao, 'SIMPLIFICADO'
    )
    if caminho_existente:
        return caminho_existente, nome_arquivo

    # ENCONTRA a hierarquia completa baseada na seleção
    dados_completos = encontrar_hierarquia_completa(dados_selecao)

    # Agora usa os dados completos para gerar o documento
    regiao = dados_completos.get('regiao', 'TODOS')
    uf = dados_completos.get('uf', 'TODOS')
    macro = dados_completos.get('macro', 'TODOS')
    regiao_saude = dados_completos.get('regiaoSaude', 'TODOS')
    municipio = dados_completos.get('municipio', 'TODOS')
    unidade = dados_completos.get('unidade', 'TODOS')

    print(f'🔄 Gerando NOVO briefing SIMPLIFICADO com hierarquia completa:')
    print(f'   Região: {regiao}, UF: {uf}, Macro: {macro}')
    print(
        f'   Região Saúde: {regiao_saude}, Município: {municipio}, Unidade: {unidade}'
    )

    # Caminho correto para a pasta static
    caminho_saida = os.path.join(
        BASE_DIR, '..', '..', '..', 'static', 'downloads', nome_arquivo
    )

    # Garante que o diretório existe
    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)

    # Cria o documento
    doc = Document()

    # Adiciona cabeçalho com logo
    adicionar_cabecalho_com_logo(doc)

    # Título
    titulo = doc.add_heading('BRIEFING SIMPLIFICADO - SISTEMA COQAE/DEEQAE', 0)
    titulo.alignment = 1  # Centralizado

    # Informações da Seleção
    doc.add_heading('DADOS DA SELEÇÃO', level=1)

    doc.add_paragraph(f'• Região: {regiao}')
    doc.add_paragraph(f'• UF: {uf}')
    doc.add_paragraph(f'• Macrorregião de Saúde: {macro}')
    doc.add_paragraph(f'• Região de Saúde: {regiao_saude}')
    doc.add_paragraph(f'• Município: {municipio}')
    doc.add_paragraph(f'• Unidade: {unidade}')

    doc.add_paragraph('_' * 50)

    # Resumo Executivo
    doc.add_heading('RESUMO EXECUTIVO', level=1)
    doc.add_paragraph(
        'Este briefing simplificado contém informações consolidadas sobre a região selecionada, com foco nos principais indicadores de saúde para tomada de decisão estratégica.'
    )

    # Principais Indicadores
    doc.add_heading('INDICADORES PRINCIPAIS - 2024/2025', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicador'
    hdr_cells[1].text = '2024'
    hdr_cells[2].text = '2025'
    hdr_cells[3].text = 'Variação'

    # Dados de exemplo com variação
    indicadores = [
        ('Estabelecimentos de Saúde', '23', '25', '+8.7%'),
        ('População Atendida', '43.210', '45.678', '+5.7%'),
        ('Procedimentos Ambulatoriais', '11.543', '12.345', '+6.9%'),
        ('Internações Realizadas', '142', '156', '+9.9%'),
        ('Cobertura Vacinal', '85%', '89%', '+4.7%'),
    ]

    for indicador, valor_2024, valor_2025, variacao in indicadores:
        row_cells = table.add_row().cells
        row_cells[0].text = indicador
        row_cells[1].text = valor_2024
        row_cells[2].text = valor_2025
        row_cells[3].text = variacao

    # Metadados
    doc.add_heading('METADADOS', level=1)
    doc.add_paragraph(
        f'• Data de Geração: {datetime.now().strftime("%d/%m/%Y às %H:%M:%S")}'
    )
    doc.add_paragraph(f'• Tipo: Briefing Simplificado')
    doc.add_paragraph(f'• Arquivo: {nome_arquivo}')
    doc.add_paragraph(f'• Sistema: COQAE/DEEQAE - Ministério da Saúde')
    doc.add_paragraph(f'• Status: Gerado agora')

    # Rodapé
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    doc.add_paragraph(
        'Sistema de Geração Automática de Briefing - COQAE/DEEQAE'
    )
    doc.add_paragraph('Ministério da Saúde - Brasil')
    doc.add_paragraph(
        f'Documento gerado automaticamente em: {datetime.now().strftime("%d/%m/%Y às %H:%M:%S")}'
    )

    # Salva o documento
    doc.save(caminho_saida)
    print(f'✅ NOVO documento simplificado salvo em: {caminho_saida}')

    # Retorna caminho relativo para o Flask
    caminho_relativo = os.path.join('static', 'downloads', nome_arquivo)
    return caminho_relativo, nome_arquivo
